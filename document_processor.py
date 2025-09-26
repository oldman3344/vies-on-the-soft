#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理模块
用于处理Excel数据提取和Word文档表格填充
"""

import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
import os
import logging
from typing import List, Dict, Any, Optional, Tuple

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器类"""
    
    def __init__(self):
        self.excel_data = None
        self.word_doc = None
        self._template_cache = {}  # 模板内容缓存
        self._last_template_path = None  # 上次使用的模板路径
    
    def set_table_borders(self, table):
        """
        为表格设置实线边框
        """
        try:
            from docx.oxml import parse_xml
            
            # 获取表格的XML元素
            tbl = table._tbl
            
            # 创建表格边框属性
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tbl.insert(0, tblPr)
            
            # 创建表格边框XML
            borders_xml = '''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tblBorders>
            '''
            
            # 解析边框XML
            tblBorders = parse_xml(borders_xml)
            
            # 移除现有的边框设置（如果有）
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            
            # 添加新的边框设置
            tblPr.append(tblBorders)
            
        except Exception as e:
            logger.warning(f"设置表格边框时出错: {e}")
            # 如果设置边框失败，继续执行，不影响主要功能
    
    def set_row_dashed_borders(self, row):
        """
        设置表格行的边框为虚线
        
        Args:
            row: Word表格行对象
        """
        try:
            for cell in row.cells:
                # 获取单元格属性
                tc_pr = cell._tc.tcPr
                if tc_pr is None:
                    tc_pr = OxmlElement('w:tcPr')
                    cell._tc.insert(0, tc_pr)
                
                # 构建虚线边框XML
                borders_xml = '''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="dashed" w:sz="4" w:space="0" w:color="000000"/>
                    <w:left w:val="dashed" w:sz="4" w:space="0" w:color="000000"/>
                    <w:bottom w:val="dashed" w:sz="4" w:space="0" w:color="000000"/>
                    <w:right w:val="dashed" w:sz="4" w:space="0" w:color="000000"/>
                </w:tcBorders>
                '''
                
                # 解析XML并添加到单元格属性中
                from docx.oxml import parse_xml
                borders_element = parse_xml(borders_xml)
                
                # 移除现有的边框设置
                existing_borders = tc_pr.find(qn('w:tcBorders'))
                if existing_borders is not None:
                    tc_pr.remove(existing_borders)
                
                # 添加新的边框设置
                tc_pr.append(borders_element)
            
            logger.info("行边框已设置为虚线")
            
        except Exception as e:
            logger.warning(f"设置行虚线边框时出错: {e}")

    def add_material_info_row(self, table, template_content=None):
        """
        在表格末尾添加申报材料说明行
        
        Args:
            table: Word表格对象
            template_content: 模板中的申报材料说明内容，如果为None则使用默认内容
        """
        try:
            # 添加新行
            new_row = table.add_row()
            
            # 合并所有单元格
            if len(new_row.cells) > 1:
                # 合并第一个单元格和其他所有单元格
                merged_cell = new_row.cells[0]
                for i in range(1, len(new_row.cells)):
                    merged_cell.merge(new_row.cells[i])
            
            # 设置文本内容 - 分为两部分设置不同样式
            cell = new_row.cells[0]
            
            # 清空默认文本
            cell.text = ""
            
            # 获取段落
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 第一部分：申报需提供以下材料： - 加大加粗黑色
            title_run = paragraph.add_run("申报需提供以下材料：")
            title_run.font.size = Pt(12)  # 加大字体
            title_run.font.name = '宋体'
            title_run.font.bold = True  # 加粗
            title_run.font.color.rgb = None  # 黑色（默认）
            
            # 第二部分：括号内容 - 红色加粗，与标题字体大小一致
            # 如果提供了模板内容，使用模板内容；否则使用默认内容
            if template_content:
                content_text = template_content
            else:
                # 默认内容（与图片中一致）
                content_text = "（下载时间：10月03号开始，请在10月05日前提供，如不及时申报，平台造成的后果smk不承任何责任）"
            
            content_run = paragraph.add_run(content_text)
            content_run.font.size = Pt(12)  # 与标题字体大小一致
            content_run.font.name = '宋体'
            content_run.font.bold = True  # 加粗
            
            # 设置红色
            from docx.shared import RGBColor
            content_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            
            # 设置虚线边框
            self.set_row_dashed_borders(new_row)
            
            logger.info("已添加申报材料说明行")
            
        except Exception as e:
            logger.warning(f"添加申报材料说明行时出错: {e}")
    
    def extract_material_info_from_template(self, word_path: str) -> Optional[str]:
        """
        从Word模板中提取申报材料说明的原始内容
        
        Args:
            word_path: Word模板文件路径
            
        Returns:
            申报材料说明内容，如果没有找到则返回None
        """
        try:
            logger.info(f"正在从模板文件提取申报材料说明: {word_path}")
            
            # 检查模板路径是否发生变化，如果变化则清除缓存
            if self._last_template_path != word_path:
                logger.info(f"检测到模板路径变化: {self._last_template_path} -> {word_path}")
                logger.info(f"清除缓存前，缓存项数: {len(self._template_cache)}")
                self._template_cache.clear()
                logger.info("✅ 因路径变化已清除缓存")
                self._last_template_path = word_path
            
            # 检查缓存
            if word_path in self._template_cache:
                cached_content = self._template_cache[word_path]
                logger.info("📋 从缓存中获取模板内容")
                logger.info(f"📋 缓存内容预览: {cached_content[:50] if cached_content else 'None'}...")
                return cached_content
            
            # 确保文件存在
            if not os.path.exists(word_path):
                logger.error(f"模板文件不存在: {word_path}")
                return None
            
            # 创建新的Document对象，确保读取最新的文件内容
            doc = Document(word_path)
            
            # 遍历所有表格，查找包含"申报需提供以下材料"的行
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if "申报需提供以下材料" in cell_text:
                            logger.info(f"在表格{table_idx}的第{row_idx}行第{cell_idx}列找到申报材料说明")
                            # 找到了申报材料说明行，提取括号内的内容
                            if "（" in cell_text and "）" in cell_text:
                                start_idx = cell_text.find("（")
                                end_idx = cell_text.find("）") + 1
                                if start_idx != -1 and end_idx != -1:
                                    template_content = cell_text[start_idx:end_idx]
                                    logger.info(f"从模板中提取到申报材料说明: {template_content}")
                                    # 存入缓存
                                    self._template_cache[word_path] = template_content
                                    # 显式清理文档对象
                                    doc = None
                                    return template_content
            
            logger.info("模板中未找到申报材料说明，将使用默认内容")
            # 存入缓存（None值也要缓存，避免重复读取）
            self._template_cache[word_path] = None
            # 显式清理文档对象
            doc = None
            return None
            
        except Exception as e:
            logger.warning(f"从模板提取申报材料说明时出错: {e}")
            return None
    
    def clear_template_cache(self):
        """
        清除模板缓存
        当用户选择新的模板文件时应该调用此方法
        """
        cache_count = len(self._template_cache)
        logger.info(f"清除模板缓存，当前缓存项数: {cache_count}")
        if cache_count > 0:
            logger.info(f"缓存的模板路径: {list(self._template_cache.keys())}")
        self._template_cache.clear()
        self._last_template_path = None
        logger.info("✅ 模板缓存已清除")
        
    def extract_excel_data(self, excel_path: str, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        从Excel文件中提取数据 - 自动检测所有有效数据行
        
        Args:
            excel_path: Excel文件路径
            sheet_name: 工作表名称，如果为None则使用第一个工作表
            
        Returns:
            提取的数据列表
        """
        try:
            # 读取Excel文件
            if sheet_name:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
            else:
                df = pd.read_excel(excel_path)
            
            logger.info(f"成功读取Excel文件: {excel_path}")
            logger.info(f"原始数据形状: {df.shape}")
            logger.info(f"列名: {list(df.columns)}")
            
            # 第一步：移除完全空白的行
            df = df.dropna(how='all')
            logger.info(f"移除空行后数据形状: {df.shape}")
            
            # 第二步：智能过滤有效数据行
            data_list = []
            for idx, row in df.iterrows():
                row_dict = row.to_dict()
                
                # 检查行是否包含有效数据
                has_valid_data = False
                non_empty_values = 0
                
                for key, value in row_dict.items():
                    if not pd.isna(value) and str(value).strip():
                        non_empty_values += 1
                        # 检查是否包含VAT相关信息或其他有效业务数据
                        value_str = str(value).lower().strip()
                        if ('vat' in value_str or 
                            len(value_str) > 2 or  # 有意义的文本
                            value_str.replace('.', '').replace(',', '').replace('-', '').isdigit()):  # 数字数据
                            has_valid_data = True
                
                # 如果行有足够的非空数据且包含有效信息，则包含此行
                if has_valid_data and non_empty_values >= 2:  # 至少2个非空字段
                    # 清理数据，将NaN替换为空字符串
                    cleaned_row = {}
                    for key, value in row_dict.items():
                        if pd.isna(value):
                            cleaned_row[key] = ""
                        else:
                            cleaned_row[key] = str(value).strip()
                    
                    data_list.append(cleaned_row)
            
            self.excel_data = data_list
            logger.info(f"✅ 自动检测并提取到 {len(data_list)} 行有效数据")
            
            # 如果没有找到数据，尝试更宽松的条件
            if not data_list:
                logger.warning("未找到符合条件的数据，尝试更宽松的筛选条件...")
                for idx, row in df.iterrows():
                    row_dict = row.to_dict()
                    non_empty_count = sum(1 for v in row_dict.values() 
                                        if not pd.isna(v) and str(v).strip())
                    
                    if non_empty_count >= 1:  # 至少1个非空字段
                        cleaned_row = {}
                        for key, value in row_dict.items():
                            if pd.isna(value):
                                cleaned_row[key] = ""
                            else:
                                cleaned_row[key] = str(value).strip()
                        data_list.append(cleaned_row)
                
                self.excel_data = data_list
                logger.info(f"📋 使用宽松条件提取到 {len(data_list)} 行数据")
            
            return data_list
            
        except Exception as e:
            logger.error(f"提取Excel数据时出错: {e}")
            raise
    
    def get_excel_sheets(self, excel_path: str) -> List[str]:
        """
        获取Excel文件中所有工作表的名称
        
        Args:
            excel_path: Excel文件路径
            
        Returns:
            工作表名称列表
        """
        try:
            excel_file = pd.ExcelFile(excel_path)
            return excel_file.sheet_names
        except Exception as e:
            logger.error(f"获取Excel工作表列表失败: {e}")
            return []

    def get_excel_columns(self, excel_path: str, sheet_name: Optional[str] = None) -> List[str]:
        """
        获取Excel文件的列名
        
        Args:
            excel_path: Excel文件路径
            sheet_name: 工作表名称，如果为None则使用第一个工作表
            
        Returns:
            列名列表
        """
        try:
            if sheet_name:
                df = pd.read_excel(excel_path, sheet_name=sheet_name, nrows=0)
            else:
                df = pd.read_excel(excel_path, nrows=0)
            return list(df.columns)
        except Exception as e:
            logger.error(f"获取Excel列名失败: {e}")
            return []
    
    def analyze_word_document(self, word_path: str) -> Dict[str, Any]:
        """
        分析Word文档结构
        
        Args:
            word_path: Word文档路径
            
        Returns:
            文档分析结果
        """
        try:
            doc = Document(word_path)
            
            analysis = {
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
                'tables_info': []
            }
            
            # 分析表格
            for i, table in enumerate(doc.tables):
                table_info = {
                    'table_index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'content': []
                }
                
                # 获取表格内容
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_info['content'].append(row_content)
                
                analysis['tables_info'].append(table_info)
            
            logger.info(f"Word文档分析完成: {analysis}")
            return analysis
            
        except Exception as e:
            logger.error(f"分析Word文档时出错: {e}")
            raise
    
    def fill_word_table(self, word_path: str, excel_data: List[Dict[str, Any]], 
                       output_path: str, table_index: int = 0,
                       column_mapping: Optional[Dict[str, str]] = None) -> str:
        """
        将Excel数据填充到Word表格中
        
        Args:
            word_path: Word模板文件路径
            excel_data: Excel数据
            output_path: 输出文件路径
            table_index: 要填充的表格索引
            column_mapping: 列映射关系 {excel_column: word_table_column}
            
        Returns:
            输出文件路径
        """
        try:
            # 首先从模板中提取申报材料说明内容
            template_material_info = self.extract_material_info_from_template(word_path)
            
            # 打开Word文档
            doc = Document(word_path)
            
            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围，文档只有 {len(doc.tables)} 个表格")

            table = doc.tables[table_index]
            
            # 如果没有提供列映射，使用默认映射
            if column_mapping is None:
                # 根据Excel数据的列名创建默认映射
                if excel_data:
                    excel_columns = list(excel_data[0].keys())
                    column_mapping = {}
                    # 获取Word表格表头
                    word_headers = [cell.text.strip() for cell in table.rows[0].cells]
                    for i, excel_col in enumerate(excel_columns):
                        if i < len(word_headers):
                            # 将Excel列名映射到Word表头
                            column_mapping[word_headers[i]] = excel_col
            
            logger.info(f"使用列映射: {column_mapping}")
            
            # 确保表格有足够的行
            current_rows = len(table.rows)
            needed_rows = len(excel_data) + 1  # +1 for header
            
            # 添加行如果需要
            while len(table.rows) < needed_rows:
                table.add_row()
            
            # 删除所有非表头行（避免合并单元格问题）
            rows_to_remove = []
            for row_idx in range(len(table.rows) - 1, 0, -1):  # 从后往前删除
                rows_to_remove.append(row_idx)
            
            for row_idx in rows_to_remove:
                table._tbl.remove(table.rows[row_idx]._tr)
            
            # 填充数据
            for row_idx, data_row in enumerate(excel_data):
                # 从第二行开始填充（第一行通常是表头）
                word_row_idx = row_idx + 1
                
                if word_row_idx >= len(table.rows):
                    table.add_row()
                
                word_row = table.rows[word_row_idx]
                
                # 填充每一列
                if column_mapping:
                    # 获取Word表格表头
                    word_headers = [cell.text.strip() for cell in table.rows[0].cells]
                    
                    # 遍历Word表格的每一列
                    for word_col_idx, word_header in enumerate(word_headers):
                        if word_header in column_mapping:
                            # 获取对应的Excel列名
                            excel_col_name = column_mapping[word_header]
                            if excel_col_name in data_row:
                                cell = word_row.cells[word_col_idx]
                                cell.text = str(data_row[excel_col_name])
                                
                                # 设置单元格对齐方式
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置表格边框为实线
            self.set_table_borders(table)
            
            # 添加申报材料说明行，使用从模板提取的内容
            self.add_material_info_row(table, template_material_info)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"Word文档已保存到: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"填充Word表格时出错: {e}")
            raise
    
    def process_documents(self, excel_path: str, word_template_path: str, 
                         output_path: str, sheet_name: Optional[str] = None,
                         column_mapping: Optional[Dict[str, str]] = None,
                         processing_mode: str = "single") -> Dict[str, Any]:
        """
        完整的文档处理流程 - 支持单个和多个处理模式
        
        Args:
            excel_path: Excel文件路径
            word_template_path: Word模板路径
            output_path: 输出文件路径
            sheet_name: Excel工作表名称，如果为None则使用第一个工作表
            column_mapping: 列映射
            processing_mode: 处理模式，"single"按公司名称分组，"multiple"按群名称分组
            
        Returns:
            包含处理结果的字典
        """
        try:
            # 1. 提取Excel数据（自动检测所有有效行）
            excel_data = self.extract_excel_data(excel_path, sheet_name)
            
            if not excel_data:
                return {
                    'success': False,
                    'error': '未找到有效的Excel数据',
                    'output_path': None,
                    'rows_filled': 0,
                    'total_rows_detected': 0
                }
            
            logger.info(f"🎯 将处理所有检测到的 {len(excel_data)} 行有效数据")
            logger.info(f"📋 处理模式: {processing_mode}")
            
            # 2. 分析Word文档
            word_analysis = self.analyze_word_document(word_template_path)
            
            # 3. 根据处理模式进行分组处理
            if processing_mode == "single":
                # 按公司名称分组，每个公司生成一个文档
                return self._process_by_company(excel_data, word_template_path, output_path, column_mapping)
            elif processing_mode == "multiple":
                # 按群名称分组，同一群生成一个文档
                return self._process_by_group(excel_data, word_template_path, output_path, column_mapping)
            else:
                # 默认处理（原有逻辑）
                return self._process_single_document(excel_data, word_template_path, output_path, column_mapping)
            
        except Exception as e:
            logger.error(f"文档处理过程中出错: {e}")
            return {
                'success': False,
                'error': str(e),
                'output_path': None,
                'rows_filled': 0,
                'total_rows_detected': 0
            }
    
    def _process_single_document(self, excel_data: List[Dict[str, Any]], 
                               word_template_path: str, output_path: str,
                               column_mapping: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        处理单个文档（原有逻辑）
        """
        try:
            # 生成输出路径（如果未提供）
            if not output_path:
                base_name = os.path.splitext(os.path.basename(word_template_path))[0]
                output_dir = os.path.dirname(word_template_path)
                output_path = os.path.join(output_dir, f"{base_name}_已填充.docx")
            
            # 填充Word表格
            result_path = self.fill_word_table(
                word_template_path, 
                excel_data, 
                output_path, 
                0, 
                column_mapping
            )
            
            logger.info(f"✅ 文档处理完成，输出文件: {result_path}")
            return {
                'success': True,
                'error': None,
                'output_path': result_path,
                'rows_filled': len(excel_data),
                'total_rows_detected': len(excel_data)
            }
        except Exception as e:
            logger.error(f"单文档处理出错: {e}")
            return {
                'success': False,
                'error': str(e),
                'output_path': None,
                'rows_filled': 0,
                'total_rows_detected': len(excel_data)
            }
    
    def _process_by_company(self, excel_data: List[Dict[str, Any]], 
                          word_template_path: str, output_path: str,
                          column_mapping: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        按公司名称分组处理，每个公司生成一个文档
        """
        try:
            # 按公司名称分组
            company_groups = {}
            company_field = '公司名称'
            
            for row in excel_data:
                company_name = row.get(company_field, '').strip()
                if not company_name:
                    company_name = '未知公司'
                
                if company_name not in company_groups:
                    company_groups[company_name] = []
                company_groups[company_name].append(row)
            
            logger.info(f"📊 按公司名称分组，共 {len(company_groups)} 个公司")
            
            # 为每个公司生成文档
            output_dir = os.path.dirname(output_path)
            base_name = os.path.splitext(os.path.basename(output_path))[0]
            generated_files = []
            total_rows = 0
            
            for company_name, company_data in company_groups.items():
                # 生成安全的文件名
                safe_company_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
                if not safe_company_name:
                    safe_company_name = "未知公司"
                
                # 智能处理括号：检查是否包含括号并相应处理
                import re
                logger.info(f"🔍 调试信息 - 原始base_name: '{base_name}'")
                logger.info(f"🔍 调试信息 - 公司名称: '{safe_company_name}'")
                
                # 检查是否包含中文括号或英文括号
                has_brackets = ('(' in base_name and ')' in base_name) or ('（' in base_name and '）' in base_name)
                logger.info(f"🔍 调试信息 - 是否包含括号: {has_brackets}")
                
                if has_brackets:
                    # 如果包含括号，替换整个括号及其内容（支持中文和英文括号）
                    # 使用正则表达式替换中文括号或英文括号及其内容
                    new_base_name = re.sub(r'[（(][^）)]*[）)]', f'({safe_company_name})', base_name)
                    logger.info(f"🔍 调试信息 - 替换后的base_name: '{new_base_name}'")
                    company_output_path = os.path.join(output_dir, f"{new_base_name}.docx")
                else:
                    # 如果不包含括号，添加括号和公司名称
                    logger.info(f"🔍 调试信息 - 不包含括号，直接添加")
                    company_output_path = os.path.join(output_dir, f"{base_name}({safe_company_name}).docx")
                
                logger.info(f"🔍 调试信息 - 最终文件路径: '{company_output_path}'")
                
                # 填充Word表格
                result_path = self.fill_word_table(
                    word_template_path, 
                    company_data, 
                    company_output_path, 
                    0, 
                    column_mapping
                )
                
                generated_files.append(result_path)
                total_rows += len(company_data)
                logger.info(f"✅ 已生成 {company_name} 的文档: {os.path.basename(result_path)} ({len(company_data)} 行数据)")
            
            return {
                'success': True,
                'error': None,
                'output_path': generated_files[0] if generated_files else None,  # 返回第一个文件路径
                'generated_files': generated_files,
                'rows_filled': total_rows,
                'total_rows_detected': len(excel_data),
                'groups_count': len(company_groups)
            }
            
        except Exception as e:
            logger.error(f"按公司分组处理出错: {e}")
            return {
                'success': False,
                'error': str(e),
                'output_path': None,
                'rows_filled': 0,
                'total_rows_detected': len(excel_data)
            }
    
    def _process_by_group(self, excel_data: List[Dict[str, Any]], 
                        word_template_path: str, output_path: str,
                        column_mapping: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        按群名称分组处理，同一群的数据生成一个文档
        """
        try:
            # 按群名称分组
            group_groups = {}
            group_field = '群名称'
            company_field = '公司名称'
            
            for row in excel_data:
                group_name = row.get(group_field, '').strip()
                if not group_name:
                    # 如果没有群名称，使用公司名称作为群名称
                    group_name = row.get(company_field, '').strip()
                    if not group_name:
                        group_name = '未知群组'
                
                if group_name not in group_groups:
                    group_groups[group_name] = []
                group_groups[group_name].append(row)
            
            logger.info(f"📊 按群名称分组，共 {len(group_groups)} 个群组")
            
            # 为每个群组生成文档
            output_dir = os.path.dirname(output_path)
            base_name = os.path.splitext(os.path.basename(output_path))[0]
            generated_files = []
            total_rows = 0
            
            for group_name, group_data in group_groups.items():
                # 获取群组中第一个公司名称用于文件名
                first_company = ""
                for row in group_data:
                    company_name = row.get(company_field, '').strip()
                    if company_name:
                        first_company = company_name
                        break
                
                if not first_company:
                    first_company = "未知公司"
                
                # 生成安全的文件名
                safe_group_name = "".join(c for c in group_name if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_company_name = "".join(c for c in first_company if c.isalnum() or c in (' ', '-', '_')).strip()
                
                if not safe_group_name:
                    safe_group_name = "未知群组"
                if not safe_company_name:
                    safe_company_name = "未知公司"
                
                # 智能处理括号：检查是否包含括号并相应处理
                import re
                # 检查是否包含中文括号或英文括号
                if ('(' in base_name and ')' in base_name) or ('（' in base_name and '）' in base_name):
                    # 如果包含括号，替换整个括号及其内容（支持中文和英文括号）
                    # 使用群组中第一个公司的名称作为文件名
                    new_base_name = re.sub(r'[（(][^）)]*[）)]', f'({safe_company_name})', base_name)
                    group_output_path = os.path.join(output_dir, f"{new_base_name}.docx")
                else:
                    # 如果不包含括号，添加括号和第一个公司名称
                    group_output_path = os.path.join(output_dir, f"{base_name}({safe_company_name}).docx")
                
                # 填充Word表格
                result_path = self.fill_word_table(
                    word_template_path, 
                    group_data, 
                    group_output_path, 
                    0, 
                    column_mapping
                )
                
                generated_files.append(result_path)
                total_rows += len(group_data)
                logger.info(f"✅ 已生成群组 {group_name} 的文档: {os.path.basename(result_path)} ({len(group_data)} 行数据)")
            
            return {
                'success': True,
                'error': None,
                'output_path': generated_files[0] if generated_files else None,  # 返回第一个文件路径
                'generated_files': generated_files,
                'rows_filled': total_rows,
                'total_rows_detected': len(excel_data),
                'groups_count': len(group_groups)
            }
            
        except Exception as e:
            logger.error(f"按群组分组处理出错: {e}")
            return {
                'success': False,
                'error': str(e),
                'output_path': None,
                'rows_filled': 0,
                'total_rows_detected': len(excel_data)
            }

def create_default_column_mapping() -> Dict[str, str]:
    """
    创建默认的列映射关系
    基于Word表格字段与Excel数据列的对应关系
    Word表格字段 -> Excel列名
    """
    return {
        '客户': '公司名称',           # 对应Excel的"公司名称"列
        '国家': '国家',             # 对应Excel的"国家"列  
        '申报方式': '申报方式',       # 对应Excel的"申报方式"列
        '申报时段': '申报时段',       # 对应Excel的"申报时段"列
        '下载数据格式': '下载数据格式',   # 对应Excel的"下载数据格式"列
        '备注(如已续费，请忽略）': '备注'  # 对应Excel的"备注"列
    }


if __name__ == "__main__":
    # 测试代码
    processor = DocumentProcessor()
    
    # 测试路径
    excel_path = "/Volumes/oldman_space/work_space/vies-on-the-soft/25.08月申报明细汇总表.xlsx"
    word_path = "/Volumes/oldman_space/work_space/vies-on-the-soft/VAT申报明细表模板.docx"
    output_path = "/Volumes/oldman_space/work_space/vies-on-the-soft/VAT申报明细表_已填充.docx"
    
    try:
        # 测试Excel数据提取
        data = processor.extract_excel_data(excel_path)
        print(f"提取到 {len(data)} 行数据")
        
        if data:
            print("前3行数据:")
            for i, row in enumerate(data[:3]):
                print(f"第{i+1}行: {row}")
        
        # 测试Word文档分析
        if os.path.exists(word_path):
            analysis = processor.analyze_word_document(word_path)
            print(f"Word文档分析结果: {analysis}")
            
            # 测试完整处理流程
            column_mapping = create_default_column_mapping()
            result = processor.process_documents(
                excel_path, word_path, output_path, 
                sheet_name=None, column_mapping=column_mapping
            )
            print(f"处理完成，输出文件: {result}")
        else:
            print(f"Word模板文件不存在: {word_path}")
            
    except Exception as e:
        print(f"测试过程中出错: {e}")